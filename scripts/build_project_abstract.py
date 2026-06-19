"""Generate VisionSafe360 filled Project Abstract .docx.

Mirrors the structure of `Project Abstract Template.docx` and fills the five
required sections with committee-ready content describing the actual
implementation. Team metadata is left as placeholders for the student to edit.
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = REPO_ROOT / "Project_Abstract_VisionSafe360.docx"

ACCENT = RGBColor(0xFF, 0x6A, 0x00)   # VisionSafe orange
DARK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)

TITLE_TEXT = (
    "VisionSafe 360: An Edge-AI Industrial Safety Intelligence Platform "
    "for Real-Time Hazard Detection, Ergonomic Risk Assessment, and "
    "Incident Lifecycle Management"
)

# --- Section bodies ---------------------------------------------------------
SECTION_1 = (
    "In industrial environments, the cost of being late is measured in "
    "lives. A worker who falls, a forklift that turns too fast, or a "
    "missing helmet can turn a normal shift into a tragedy in seconds. "
    "Yet most workplaces still rely on cameras that only record what has "
    "already happened — useful for investigation, useless for prevention."
    "\n\n"
    "VisionSafe 360 turns this passive infrastructure into an active line "
    "of defense. It is an AI-powered industrial safety platform that uses "
    "the cameras a factory, warehouse, or construction site already owns "
    "to recognize danger as it unfolds — and to intervene before the "
    "accident occurs."
    "\n\n"
    "The platform brings together, in a single unified solution, the four "
    "risks most responsible for industrial injuries: PPE non-compliance, "
    "unsafe forklift-worker proximity, falls and fall-from-height events, "
    "and chronic ergonomic strain in repetitive tasks. Each detected "
    "hazard is validated to suppress false alarms, escalated through a "
    "real-time command center, and delivered as a mobile alert to the "
    "responsible supervisor with severity, location, and video evidence "
    "attached."
    "\n\n"
    "This is the project's unique value proposition. VisionSafe 360 is "
    "not another surveillance tool and not another single-purpose AI "
    "demonstration. It is the first unified safety ecosystem that "
    "combines PPE compliance monitoring, forklift-worker risk detection, "
    "fall detection, ergonomic assessment, mobile incident response, and "
    "real-time analytics inside one integrated platform — replacing "
    "fragmented tools and reactive workflows with continuous, "
    "evidence-driven accident prevention."
)

SECTION_2 = (
    "VisionSafe 360 is designed for adoption, not just admiration. Its "
    "single greatest practical strength is that organizations do not need "
    "to replace anything to benefit from it: the platform runs on the "
    "CCTV cameras they already operate, removing the cost and disruption "
    "barriers that have historically blocked AI safety adoption in "
    "industrial settings."
    "\n\n"
    "All analysis runs locally at the edge, so video never leaves the "
    "facility. This eliminates continuous cloud-upload costs, satisfies "
    "the data-sovereignty and privacy requirements that block many "
    "industrial customers from adopting AI safety solutions, and "
    "guarantees real-time response without dependence on network speed."
    "\n\n"
    "Scalability is built into the architecture. A small facility can "
    "start with one edge device and a handful of cameras and grow into a "
    "multi-site enterprise deployment on the same platform, with cost "
    "scaling linearly per camera. Each site can also enable only the "
    "safety capabilities relevant to its operations, controlling cost "
    "without sacrificing capability."
    "\n\n"
    "For organizations, the business case is direct. Every prevented "
    "fall, collision, or PPE violation is a measurable reduction in "
    "injuries, regulatory exposure, insurance premiums, and lost "
    "productivity. Aligned with OSHA, NIOSH, and ISO 45001 trends and the "
    "global shift toward continuous safety monitoring, the platform meets "
    "a market that is actively looking for it."
)

SECTION_3_INTRO = (
    "The completed system has been implemented, integrated, and "
    "demonstrated as a fully working industrial safety platform. Its "
    "capabilities are best understood as four cohesive layers of an "
    "integrated solution rather than a list of independent features:"
)

SECTION_3_GROUPS = [
    ("Unified AI safety intelligence",
     "A single platform continuously monitors PPE compliance, "
     "forklift-worker proximity, fall events, and ergonomic posture using "
     "the internationally recognized RULA and REBA assessment "
     "frameworks. Unifying these four hazard categories — historically "
     "scattered across separate vendor tools — is the central technical "
     "contribution of the project."),

    ("Real-time command and response center",
     "A web dashboard provides multi-camera live monitoring, alert "
     "triage, configurable safety zones, full incident lifecycle "
     "management, analytics, and exportable safety reports. A mobile "
     "push channel delivers prioritized alerts to supervisors with "
     "severity, location, and video evidence attached, enabling "
     "immediate intervention from anywhere on site."),

    ("Intelligent decision support",
     "Detections are stabilized and validated before triggering alerts, "
     "virtually eliminating alert fatigue. Analytics expose recurring "
     "hazard patterns by zone, shift, or activity, transforming raw "
     "incidents into actionable insight for training, engineering "
     "controls, and management reporting."),

    ("Resilient and accessible deployment",
     "The platform continues detecting and alerting even when network "
     "connectivity is interrupted, with no incident lost and automatic "
     "synchronization once connectivity returns. A bilingual interface "
     "(English and Arabic) and full accessibility compliance allow "
     "adoption across diverse industrial workforces and supervisors of "
     "varying technical backgrounds."),
]

SECTION_3_CLOSE = (
    "Together, these results demonstrate a deployable, operationally "
    "complete platform — not a research prototype."
)

SECTION_4_INTRO = (
    "VisionSafe 360 changes how safety is practiced on the factory floor. "
    "Instead of reviewing recordings after an injury, supervisors are "
    "alerted within seconds — while the hazard is still preventable. "
    "Instead of relying on periodic walkthroughs, every camera becomes a "
    "tireless inspector applying consistent judgment across shifts, "
    "sites, and languages. The benefits span four dimensions:"
)

SECTION_4_GROUPS = [
    ("Operational impact",
     "Response time collapses from minutes or hours to seconds. For "
     "falls in particular — where every minute of delay can be "
     "life-threatening — this is the difference between recovery and "
     "tragedy. Continuous PPE and proximity monitoring address the most "
     "common causes of struck-by injuries and head trauma in industrial "
     "environments."),

    ("Financial impact",
     "The platform turns hidden injury costs — lost-time incidents, "
     "regulatory penalties, insurance premiums, and litigation — into "
     "preventable expenses. Always-on ergonomic monitoring captures "
     "musculoskeletal risk early, when corrective action is dramatically "
     "cheaper than treatment."),

    ("Organizational impact",
     "Fragmented spreadsheets and informal reporting are replaced by a "
     "structured incident lifecycle. Every alert is recorded with "
     "evidence, routed to the responsible supervisor, acknowledged, and "
     "resolved with a complete audit trail. Analytics expose recurring "
     "hazard patterns, giving managers a basis for targeted training and "
     "engineering controls instead of guesswork."),

    ("Cultural impact",
     "The platform enables a transition from blame-based safety to "
     "evidence-based safety. Workers benefit from transparent, objective "
     "monitoring; supervisors gain the data to act with confidence; and "
     "management gains both the assurance that the safety program is "
     "working and the proof to demonstrate it to regulators, partners, "
     "and customers."),
]

SECTION_5_INTRO = (
    "VisionSafe 360 is built for any environment where workers, machines, "
    "and hazards share space. Its most direct beneficiaries include:"
)

SECTION_5_BULLETS = [
    ("Manufacturing facilities and assembly plants",
     "continuous PPE enforcement, real-time ergonomic monitoring on "
     "production lines, and immediate fall response — reducing "
     "injury-driven downtime and supporting ISO 45001 and OSHA "
     "compliance."),

    ("Warehouses and logistics centers",
     "proximity intelligence that prevents one of the deadliest "
     "industrial hazards, with configurable safety zones automatically "
     "enforcing pedestrian and vehicle corridors."),

    ("Construction sites",
     "helmet and vest compliance, fall detection, and zone-based "
     "supervision in environments where conventional oversight cannot "
     "scale."),

    ("Oil, gas, and heavy industry",
     "restricted-zone monitoring, emergency-exit oversight, and "
     "continuous PPE compliance in safety-critical areas where a single "
     "violation can be catastrophic."),

    ("Ports, mining, and quarrying operations",
     "vehicle and pedestrian intelligence in high-traffic environments "
     "where heavy machinery and ground workers must operate side by "
     "side."),

    ("Safety officers and HSE managers",
     "a single command center that replaces fragmented reporting tools, "
     "freeing time previously consumed by manual audits and paperwork."),

    ("Floor supervisors",
     "prioritized, evidence-rich mobile alerts that enable informed "
     "intervention exactly when and where it is needed."),

    ("Workers themselves",
     "the strongest beneficiaries — faster emergency response, "
     "protection against the most common workplace hazards, ergonomic "
     "interventions before chronic injuries develop, and an objective "
     "safety culture that protects them rather than blames them."),

    ("Regulators, auditors, and insurance providers",
     "immutable, evidence-backed safety telemetry that supports "
     "compliance verification, fair risk pricing, and continuous "
     "oversight."),
]


# --- Styling helpers --------------------------------------------------------
def _set_run(run, *, size=None, bold=False, color=None, font_name="Calibri"):
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_horizontal_line(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "FF6A00")
    pBdr.append(bottom)
    p_pr.append(pBdr)


def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{number}. {title}")
    _set_run(run, size=14, bold=True, color=ACCENT)
    _add_horizontal_line(p)
    return p


def add_body_paragraph(doc, text, *, italic=False, bullet=False):
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
    else:
        p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    _set_run(run, size=11, color=DARK)
    run.italic = italic
    return p


def add_multi_paragraph(doc, text):
    for chunk in text.split("\n\n"):
        add_body_paragraph(doc, chunk.strip())


def add_labeled_bullet(doc, label, body):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(f"{label} — ")
    _set_run(run, size=11, bold=True, color=DARK)
    run2 = p.add_run(body)
    _set_run(run2, size=11, color=DARK)
    return p


# --- Document construction --------------------------------------------------
def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Header: Faculty / University
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("FACULTY OF ARTIFICIAL INTELLIGENCE")
    _set_run(run, size=14, bold=True, color=ACCENT)
    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = header2.add_run("Egyptian Russian University")
    _set_run(run2, size=12, bold=True, color=MUTED)

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("PROJECT ABSTRACT")
    _set_run(run, size=18, bold=True, color=DARK)

    # Project title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Project Title:")
    _set_run(run, size=11, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(TITLE_TEXT)
    _set_run(run, size=13, bold=True, color=DARK)

    # Team / Supervisor / Year table
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    headers = ["Student Team", "Supervisor", "Academic Year"]
    contents = [
        "1. [Student Name]\n2. [Student Name]\n3. [Student Name]\n4. [Student Name]",
        "Prof./Dr. [Supervisor Name]",
        "20XX / 20XX",
    ]
    for col, text in enumerate(headers):
        cell = table.rows[0].cells[col]
        _shade_cell(cell, "FF6A00")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        _set_run(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for col, text in enumerate(contents):
        cell = table.rows[1].cells[col]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for i, line in enumerate(text.split("\n")):
            para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line)
            _set_run(run, size=10.5, color=DARK)

    doc.add_paragraph()  # spacer

    # Section 1
    add_section_heading(doc, 1, "Project Idea / Abstract")
    add_multi_paragraph(doc, SECTION_1)

    # Section 2
    add_section_heading(doc, 2, "Project Feasibility / Value")
    add_multi_paragraph(doc, SECTION_2)

    # Section 3
    add_section_heading(doc, 3, "Project Results / Outputs")
    add_body_paragraph(doc, SECTION_3_INTRO)
    for label, body in SECTION_3_GROUPS:
        add_labeled_bullet(doc, label, body)
    add_body_paragraph(doc, SECTION_3_CLOSE)

    # Section 4
    add_section_heading(doc, 4, "Benefits and Utilization")
    add_body_paragraph(doc, SECTION_4_INTRO)
    for label, body in SECTION_4_GROUPS:
        add_labeled_bullet(doc, label, body)

    # Section 5
    add_section_heading(doc, 5, "Target Beneficiaries")
    add_body_paragraph(doc, SECTION_5_INTRO)
    for label, body in SECTION_5_BULLETS:
        add_labeled_bullet(doc, label, body)

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
